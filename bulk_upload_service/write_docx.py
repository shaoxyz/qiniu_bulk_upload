import redis
from docx import Document
from docx.shared import Inches

def read_from_redis():
    redisClient = redis.Redis(decode_responses=True)
    data = redisClient.hgetall('qiniu_bulk_upload')
    return data

def write_docx(data):
    document = Document()
    document.add_heading('汇总', 0)
    for k, v in data.items():
        document.add_paragraph(k)

        document.add_paragraph(v)

    document.save('汇总.docx')

    print("Done")

if __name__ == '__main__':
    data = read_from_redis()
    write_docx(data)